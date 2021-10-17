import json
from os import listdir
from os.path import getsize

import docx

from extract import DialogFlowAPI


class AddressParser:
    def __init__(self, docs_path: str = None, format_: str = 'docx', ai_project_path: str = None):
        self.DOCS_PATH = docs_path
        self.DOCS = []  # variable for list of docs path
        self.FORMAT = None
        if ai_project_path:
            self.AI = DialogFlowAPI(
                project_id=json.load(open(ai_project_path, 'r')).get('project_id'),
                creds_path=ai_project_path)
        self.AI_EXTRACT_VALUES = [
            'postcode',
            'city',
            'state'
        ]
        self.CONTRACT_KEYWORDS = [
            'договор',
            'договора',
            'договоры',

        ]
        self.SERVICE_KEYWORDS = [
            'место оказания',
            'место нахождения',
            'место поставки',
            'оказания услуг',
            'место проведения работ',
            'до склада покупателя',
            'получатель услуг',
            'получатель товара',
            'адрес грузополучателя',
            'грузополучатель',
            'место выполнения'

        ]

        if not self.DOCS_PATH:
            return
        if format_:
            self.FORMAT = format_

        if self.FORMAT:
            for doc in listdir(docs_path):
                if doc[-len(self.FORMAT):] == self.FORMAT:  # this is our format
                    self.DOCS.append(doc)
        else:
            self.DOCS = listdir(docs_path)

    def get_contract_docs(self):
        '''
        :return: File with contains contract
        '''

        """If a document has keywords about contract in name, we need to define that document as contract"""
        return [self.DOCS_PATH + doc for doc in self.DOCS]

        for doc in self.DOCS:
            for keyword in self.CONTRACT_KEYWORDS:
                if keyword in doc.lower():
                    return self.DOCS_PATH + doc

        """If in path we have 2 documents, we need to return a biggest document"""

        if len(self.DOCS) <= 2:
            biggest_doc = {
                'path': '',
                'size': -1
            }
            for doc in self.DOCS:
                new_size = getsize(self.DOCS_PATH + doc)
                if new_size > biggest_doc['size']:
                    biggest_doc['path'] = self.DOCS_PATH + doc
                    biggest_doc['size'] = new_size
            return biggest_doc['path']

    def extract_place_of_service(self, doc_path: str):
        try:
            doc = docx.Document(doc_path)
        except docx.opc.exceptions.PackageNotFoundError:
            return
        for index, par in enumerate(doc.paragraphs):
            text = None
            if [k_wd for k_wd in self.SERVICE_KEYWORDS if k_wd in par.text.lower()]:  # TODO: Make better
                text = par.text
                for k_wd in self.SERVICE_KEYWORDS:
                    text = text.lower().replace(k_wd, '')
                try:
                    ai_return = self.AI.extract(text)
                except:
                    pass
                else:
                    return {
                        value: ai_return.get(value) for value in self.AI_EXTRACT_VALUES
                    }

    def get_all_addresses(self):
        addresses = []
        for doc in self.DOCS:
            doc_path = self.DOCS_PATH + doc
            document = docx.Document(doc_path)
            for index, par in enumerate(document.paragraphs):
                if 'г.' in par.text or 'ул.' in par.text or 'область' in par.text:
                    adrs_txt = ''
                    try:
                        adrs_txt += document.paragraphs[index - 1].text
                    except IndexError:
                        ...
                    try:
                        adrs_txt += document.paragraphs[index].text
                    except IndexError:
                        ...
                    try:
                        adrs_txt += document.paragraphs[index + 1].text
                    except IndexError:
                        ...
                    addresses.append(adrs_txt)
        return addresses
